"""Local pre-processing of inputs.

Everything in this module runs locally (no AI). It turns a PDF, Word document,
or URL into a cleaned, normalized ``SourceDocument`` that is then handed to the
AI layer for analysis.
"""
from __future__ import annotations

import ipaddress
import re
import socket
from pathlib import Path
from urllib.parse import urlparse

from .models import SourceDocument

# Restrict URL ingestion to trusted job-hosting domains to prevent SSRF.
# Extend this set only with domains your organization explicitly trusts.
ALLOWED_JOB_URL_HOSTS = {
    "company.com",
    "www.company.com",
    "careers.company.com",
}


def _is_allowed_host(hostname: str) -> bool:
    host = hostname.lower().rstrip(".")
    for allowed in ALLOWED_JOB_URL_HOSTS:
        allowed_host = allowed.lower().rstrip(".")
        if host == allowed_host or host.endswith(f".{allowed_host}"):
            return True
    return False

# Section headers commonly seen in job descriptions. Used for a light-touch
# heuristic split so the AI gets some structure to work with.
SECTION_PATTERNS = {
    "summary": r"(?:job\s+summary|role\s+summary|about\s+the\s+role|overview|position\s+summary)",
    "responsibilities": r"(?:responsibilities|duties|what\s+you'?ll\s+do|key\s+accountabilities|the\s+role)",
    "requirements": r"(?:requirements|qualifications|what\s+you'?ll\s+need|skills|experience\s+required|must\s+have)",
    "nice_to_have": r"(?:nice\s+to\s+have|preferred|bonus|desirable)",
}


def clean_text(text: str) -> str:
    """Normalize whitespace and strip junk while preserving line structure."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    # collapse 3+ blank lines down to a single blank line
    text = re.sub(r"\n\s*\n\s*\n+", "\n\n", text)
    lines = [ln.strip() for ln in text.split("\n")]
    return "\n".join(lines).strip()


def guess_role_title(text: str, source_name: str = "") -> str:
    """Best-effort job title extraction from the first lines / filename."""
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if 3 <= len(line) <= 90 and not line.endswith((".", ":")):
            return line
        break
    stem = Path(source_name).stem.replace("_", " ").replace("-", " ").strip()
    return stem.title() if stem else "Unknown Role"


def split_sections(text: str) -> dict[str, str]:
    """Heuristically split text into known sections by header keywords."""
    lowered = text.lower()
    hits: list[tuple[int, str]] = []
    for name, pat in SECTION_PATTERNS.items():
        m = re.search(rf"(?m)^\s*{pat}\b.*$", lowered)
        if m:
            hits.append((m.start(), name))
    hits.sort()
    sections: dict[str, str] = {}
    for i, (start, name) in enumerate(hits):
        end = hits[i + 1][0] if i + 1 < len(hits) else len(text)
        chunk = text[start:end].strip()
        chunk = re.sub(r"^[^\n]*\n", "", chunk, count=1)  # drop the header line
        sections[name] = chunk.strip()
    return sections


# ---------------------------------------------------------------------------
# Extractors
# ---------------------------------------------------------------------------

def extract_pdf(path: str | Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts = [(page.extract_text() or "") for page in reader.pages]
    return "\n".join(parts)


def extract_docx(path: str | Path) -> str:
    import docx  # python-docx

    document = docx.Document(str(path))
    parts: list[str] = [p.text for p in document.paragraphs]
    # include simple table content
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _validate_public_http_url(url: str) -> str:
    parsed = urlparse(url)
    if parsed.scheme not in {"http", "https"}:
        raise ValueError("URL must start with http:// or https://")
    if not parsed.hostname:
        raise ValueError("URL must include a valid hostname")
    if parsed.username or parsed.password:
        raise ValueError("URLs with embedded credentials are not allowed")
    if not _is_allowed_host(parsed.hostname):
        raise ValueError("URL hostname is not in the allowed list")

    try:
        infos = socket.getaddrinfo(parsed.hostname, None)
    except socket.gaierror as exc:
        raise ValueError("Could not resolve URL hostname") from exc

    for info in infos:
        ip_text = info[4][0]
        ip = ipaddress.ip_address(ip_text)
        if (
            ip.is_private
            or ip.is_loopback
            or ip.is_link_local
            or ip.is_multicast
            or ip.is_reserved
            or ip.is_unspecified
        ):
            raise ValueError("URL resolves to a non-public IP address, which is not allowed")

    return parsed.geturl()


def extract_url(url: str, timeout: int = 20) -> str:
    import requests
    from bs4 import BeautifulSoup

    safe_url = _validate_public_http_url(url)
    headers = {"User-Agent": "Mozilla/5.0 (compatible; AIOpportunityFinder/1.0)"}
    resp = requests.get(safe_url, headers=headers, timeout=timeout)
    resp.raise_for_status()
    soup = BeautifulSoup(resp.text, "lxml")
    for tag in soup(["script", "style", "noscript", "header", "footer", "nav", "svg"]):
        tag.decompose()
    # prefer the most content-dense container if present
    main = soup.find("main") or soup.find("article") or soup.body or soup
    return main.get_text(separator="\n")


# ---------------------------------------------------------------------------
# Public entry points
# ---------------------------------------------------------------------------

def _build_document(raw: str, source_type: str, source_name: str) -> SourceDocument:
    cleaned = clean_text(raw)
    if not cleaned:
        raise ValueError("No readable text could be extracted from the input.")
    doc = SourceDocument(
        source_type=source_type,
        source_name=source_name,
        role_title=guess_role_title(cleaned, source_name),
        raw_text=cleaned,
        sections=split_sections(cleaned),
        char_count=len(cleaned),
    )
    return doc


def ingest_file(path: str | Path, source_name: str | None = None) -> SourceDocument:
    path = Path(path)
    name = source_name or path.name
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        raw, stype = extract_pdf(path), "pdf"
    elif suffix in (".docx", ".doc"):
        raw, stype = extract_docx(path), "docx"
    elif suffix in (".txt", ".md"):
        raw, stype = path.read_text(encoding="utf-8", errors="ignore"), "text"
    else:
        raise ValueError(f"Unsupported file type: {suffix}")
    return _build_document(raw, stype, name)


def ingest_url(url: str) -> SourceDocument:
    raw = extract_url(url)
    return _build_document(raw, "url", url)


def ingest_text(text: str, source_name: str = "Pasted text") -> SourceDocument:
    return _build_document(text, "text", source_name)
